"""
docx_preserver.py
-----------------
Update DOCX content while preserving the original formatting, fonts,
layout, alignment, and design.

Strategy:
  1. Open the original DOCX with python-docx
  2. For each section that has improved content:
     a. Identify the paragraph range (content_para_start → content_para_end)
     b. Clear the text of those paragraphs while preserving formatting
     c. Fill in the improved text, distributing across existing paragraphs
     d. If the improved text has more/fewer paragraphs, add/remove as needed
        while cloning formatting from adjacent paragraphs
  3. Save as a new file (never modify original)
"""

import copy
import io
import logging
import re
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


def _clone_paragraph_format(source_para, target_para):
    """Copy paragraph formatting (alignment, spacing, indent) from source to target."""
    try:
        sf = source_para.paragraph_format
        tf = target_para.paragraph_format
        tf.alignment = sf.alignment
        tf.left_indent = sf.left_indent
        tf.right_indent = sf.right_indent
        tf.first_line_indent = sf.first_line_indent
        tf.space_before = sf.space_before
        tf.space_after = sf.space_after
        tf.line_spacing = sf.line_spacing
        tf.line_spacing_rule = sf.line_spacing_rule
    except Exception:
        pass


def _clone_run_format(source_run, target_run):
    """Copy character formatting (font, size, bold, italic, color) from source to target."""
    try:
        sf = source_run.font
        tf = target_run.font
        tf.name = sf.name
        tf.size = sf.size
        tf.bold = sf.bold
        tf.italic = sf.italic
        tf.underline = sf.underline
        if sf.color and sf.color.rgb:
            tf.color.rgb = sf.color.rgb
    except Exception:
        pass


def _get_reference_run_format(para):
    """Get formatting from the first non-empty run in a paragraph, or any run."""
    for run in para.runs:
        if run.text.strip():
            return run
    if para.runs:
        return para.runs[0]
    return None


def _clear_paragraph_text(para):
    """Clear all text from a paragraph while preserving the first run's formatting."""
    for run in para.runs:
        run.text = ''


def _set_paragraph_text(para, text: str):
    """Set paragraph text, preserving formatting of the first run."""
    if para.runs:
        # Clear all runs except the first, set text on first
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ''
    else:
        # No runs exist — add one
        para.add_run(text)


def update_docx_in_place(
    filepath: str,
    section_mappings: List[Dict],
    accepted_sections: Optional[List[str]] = None,
) -> io.BytesIO:
    """
    Apply improved content to a DOCX file while preserving formatting.

    Args:
        filepath: Path to the original DOCX file
        section_mappings: List of mapping dicts from section_mapper.map_sections()
        accepted_sections: List of section IDs to apply (None = apply all changes)

    Returns:
        BytesIO buffer containing the updated DOCX
    """
    import docx as docx_lib
    from lxml import etree

    doc = docx_lib.Document(filepath)
    paragraphs = doc.paragraphs

    # Process sections in reverse order to not invalidate indices
    changes = []
    for mapping in section_mappings:
        section_id = mapping.get('original_id', '')

        # Skip if not accepted
        if accepted_sections is not None and section_id not in accepted_sections:
            continue

        # Skip if no changes
        if not mapping.get('has_changes'):
            continue

        para_start = mapping.get('content_para_start')
        para_end = mapping.get('content_para_end')
        improved = mapping.get('improved_content', '')

        if para_start is None or para_end is None:
            continue

        changes.append({
            'section_id': section_id,
            'para_start': para_start,
            'para_end': para_end,
            'improved': improved,
        })

    # Sort by para_start descending so we process from bottom to top
    changes.sort(key=lambda c: c['para_start'], reverse=True)

    for change in changes:
        para_start = change['para_start']
        para_end = min(change['para_end'], len(paragraphs))
        improved_text = change['improved']

        if para_start >= len(paragraphs):
            continue

        # Split improved text into lines
        improved_lines = [l for l in improved_text.split('\n')]
        # Remove empty trailing lines
        while improved_lines and not improved_lines[-1].strip():
            improved_lines.pop()

        original_para_count = para_end - para_start
        new_line_count = len(improved_lines)

        # Get reference paragraph for formatting
        ref_para = paragraphs[para_start] if para_start < len(paragraphs) else None

        if new_line_count <= original_para_count:
            # We have enough paragraphs — fill them in
            for i, line in enumerate(improved_lines):
                idx = para_start + i
                if idx < len(paragraphs):
                    _set_paragraph_text(paragraphs[idx], line)

            # Clear any extra paragraphs (don't remove to preserve structure)
            for i in range(new_line_count, original_para_count):
                idx = para_start + i
                if idx < len(paragraphs):
                    _clear_paragraph_text(paragraphs[idx])
                    _set_paragraph_text(paragraphs[idx], '')
        else:
            # Fill existing paragraphs
            for i in range(original_para_count):
                idx = para_start + i
                if idx < len(paragraphs) and i < new_line_count:
                    _set_paragraph_text(paragraphs[idx], improved_lines[i])

            # We need more paragraphs — insert after the last existing content paragraph
            insert_after_idx = para_end - 1
            if insert_after_idx < 0:
                insert_after_idx = 0

            extra_lines = improved_lines[original_para_count:]
            # Insert new paragraphs by manipulating the XML
            if insert_after_idx < len(paragraphs) and extra_lines:
                ref_element = paragraphs[insert_after_idx]._element
                parent = ref_element.getparent()

                for extra_line in reversed(extra_lines):
                    # Create new paragraph element cloning format
                    new_p = copy.deepcopy(ref_element)
                    # Clear runs in the clone
                    for r in new_p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                        new_p.remove(r)
                    # Add a run with the new text
                    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    new_r = etree.SubElement(new_p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')

                    # Clone run properties from reference if available
                    ref_run = _get_reference_run_format(paragraphs[insert_after_idx])
                    if ref_run and ref_run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr') is not None:
                        rpr = copy.deepcopy(ref_run._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'))
                        new_r.insert(0, rpr)

                    new_t = etree.SubElement(new_r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    new_t.text = extra_line
                    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

                    ref_element.addnext(new_p)

        logger.info('Updated section "%s" (%d → %d lines)',
                    change['section_id'], original_para_count, new_line_count)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
