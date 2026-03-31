"""
section_parser.py
-----------------
Parse a resume (plain text, DOCX, or PDF) into structured sections,
preserving position/index information so content can be replaced in-place.

Returns a list of Section dicts:
  {
    "id": "experience",
    "heading": "Work Experience",
    "heading_line": 12,         # 0-based line index of the heading
    "content_start": 13,        # first content line
    "content_end": 28,          # exclusive end line
    "content": "...",           # raw text of the section body
    "raw_lines": [...]          # list of original lines
  }
"""

import re
from typing import List, Dict, Optional

# ── Heading patterns (case-insensitive) mapped to canonical section IDs ──
SECTION_PATTERNS = {
    'personal_info': [
        r'^personal\s*(info(rmation)?|details?)?\s*:?\s*$',
        r'^contact\s*(info(rmation)?|details?)?\s*:?\s*$',
    ],
    'summary': [
        r'^(professional\s+)?summary\s*:?\s*$',
        r'^(career\s+)?objective\s*:?\s*$',
        r'^(career\s+|professional\s+)?profile\s*:?\s*$',
        r'^about\s*(me)?\s*:?\s*$',
        r'^executive\s+summary\s*:?\s*$',
    ],
    'experience': [
        r'^(work\s+|professional\s+)?experience\s*:?\s*$',
        r'^employment\s*(history)?\s*:?\s*$',
        r'^(career|work)\s+history\s*:?\s*$',
    ],
    'education': [
        r'^education(al)?\s*(background|qualifications?)?\s*:?\s*$',
        r'^academic\s*(background|history|qualifications?)?\s*:?\s*$',
        r'^education\s*(and|&)\s*training\s*:?\s*$',
    ],
    'skills': [
        r'^(key\s+|core\s+|technical\s+)?skills\s*(&\s*(abilities|competencies))?\s*:?\s*$',
        r'^(technical\s+)?proficienc(y|ies)\s*:?\s*$',
        r'^tools?\s*(&|and)\s*technolog(y|ies)\s*:?\s*$',
        r'^competenc(y|ies)\s*:?\s*$',
    ],
    'projects': [
        r'^(personal\s+|selected\s+|notable\s+|featured\s+)?projects?\s*:?\s*$',
        r'^portfolio\s*:?\s*$',
        r'^case\s+stud(y|ies)\s*:?\s*$',
    ],
    'certifications': [
        r'^certific?ations?\s*:?\s*$',
        r'^licens(e|es)\s*(&|and)?\s*certific?ations?\s*:?\s*$',
        r'^credentials?\s*:?\s*$',
    ],
    'achievements': [
        r'^achievements?\s*:?\s*$',
        r'^accomplishments?\s*:?\s*$',
        r'^awards?\s*((&|and)\s*honors?)?\s*:?\s*$',
        r'^honors?\s*:?\s*$',
    ],
    'languages': [
        r'^languages?\s*:?\s*$',
    ],
    'interests': [
        r'^(hobbies?\s*(&|and)\s*)?interests?\s*:?\s*$',
        r'^hobbies?\s*:?\s*$',
    ],
    'references': [
        r'^references?\s*:?\s*$',
    ],
}

# Compiled patterns
_COMPILED = {}
for section_id, patterns in SECTION_PATTERNS.items():
    _COMPILED[section_id] = [re.compile(p, re.IGNORECASE) for p in patterns]


def _is_heading_line(line: str) -> bool:
    """Heuristic: line looks like a section heading (short, possibly ALL-CAPS or title-case, no bullets)."""
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return False
    # Skip bullet lines
    if re.match(r'^[\-•\*\d]', stripped):
        return False
    # ALL-CAPS line with mostly letters (>= 3 letters)
    alpha = re.sub(r'[^A-Za-z]', '', stripped)
    if len(alpha) >= 3 and stripped == stripped.upper():
        return True
    # Title-case-ish (first word capitalized, short)
    if len(stripped.split()) <= 6 and stripped[0].isupper():
        # Check against known patterns
        for section_id, compiled_list in _COMPILED.items():
            for pat in compiled_list:
                if pat.match(stripped):
                    return True
    return False


def _match_heading(line: str) -> Optional[str]:
    """Return canonical section ID if line matches a known heading, else None."""
    stripped = line.strip().rstrip(':').strip()
    if not stripped:
        return None
    for section_id, compiled_list in _COMPILED.items():
        for pat in compiled_list:
            if pat.match(stripped):
                return section_id
    # Fallback: ALL-CAPS heading that didn't match known patterns
    alpha = re.sub(r'[^A-Za-z ]', '', stripped).strip()
    if len(alpha) >= 3 and stripped == stripped.upper():
        # Try fuzzy match against section names
        lower = alpha.lower()
        for section_id in SECTION_PATTERNS:
            if section_id in lower or lower in section_id:
                return section_id
        return 'other_' + re.sub(r'\W+', '_', lower)[:30]
    return None


def parse_sections_from_text(text: str) -> List[Dict]:
    """
    Parse resume text into structured sections.
    Returns list of section dicts with position info.
    """
    lines = text.split('\n')
    sections: List[Dict] = []
    current_id = None
    current_heading = ''
    current_heading_line = 0
    content_lines: List[str] = []
    preamble_lines: List[str] = []

    for i, line in enumerate(lines):
        matched_id = _match_heading(line)
        if matched_id is not None:
            # Save previous section
            if current_id:
                sections.append({
                    'id': current_id,
                    'heading': current_heading,
                    'heading_line': current_heading_line,
                    'content_start': current_heading_line + 1,
                    'content_end': i,
                    'content': '\n'.join(content_lines).strip(),
                    'raw_lines': content_lines[:],
                })
            elif preamble_lines:
                # Text before the first heading is treated as personal_info/header
                sections.append({
                    'id': 'personal_info',
                    'heading': '',
                    'heading_line': 0,
                    'content_start': 0,
                    'content_end': i,
                    'content': '\n'.join(preamble_lines).strip(),
                    'raw_lines': preamble_lines[:],
                })

            current_id = matched_id
            current_heading = line.strip()
            current_heading_line = i
            content_lines = []
        else:
            if current_id:
                content_lines.append(line)
            else:
                preamble_lines.append(line)

    # Don't forget the last section
    if current_id:
        sections.append({
            'id': current_id,
            'heading': current_heading,
            'heading_line': current_heading_line,
            'content_start': current_heading_line + 1,
            'content_end': len(lines),
            'content': '\n'.join(content_lines).strip(),
            'raw_lines': content_lines[:],
        })
    elif preamble_lines and not sections:
        # Entire resume had no recognizable headings
        sections.append({
            'id': 'full_resume',
            'heading': '',
            'heading_line': 0,
            'content_start': 0,
            'content_end': len(lines),
            'content': '\n'.join(preamble_lines).strip(),
            'raw_lines': preamble_lines[:],
        })

    return sections


def parse_sections_from_docx(filepath: str) -> List[Dict]:
    """
    Parse DOCX file into structured sections, preserving paragraph indices.
    Each section's 'para_start' and 'para_end' reference python-docx paragraph indices.
    """
    import docx as docx_lib
    doc = docx_lib.Document(filepath)

    sections: List[Dict] = []
    current_id = None
    current_heading = ''
    current_heading_para = 0
    content_parts: List[str] = []
    preamble_parts: List[str] = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Check if paragraph uses a Heading style
        is_style_heading = para.style and para.style.name and 'Heading' in para.style.name
        matched_id = _match_heading(text) if text else None

        if matched_id is not None or (is_style_heading and text):
            # If we have a style heading but no pattern match, derive an ID
            if matched_id is None and is_style_heading:
                slug = re.sub(r'\W+', '_', text.lower())[:30]
                matched_id = slug if slug else 'section'

            if current_id:
                sections.append({
                    'id': current_id,
                    'heading': current_heading,
                    'para_start': current_heading_para,
                    'content_para_start': current_heading_para + 1,
                    'content_para_end': idx,
                    'content': '\n'.join(content_parts).strip(),
                })
            elif preamble_parts:
                sections.append({
                    'id': 'personal_info',
                    'heading': '',
                    'para_start': 0,
                    'content_para_start': 0,
                    'content_para_end': idx,
                    'content': '\n'.join(preamble_parts).strip(),
                })

            current_id = matched_id
            current_heading = text
            current_heading_para = idx
            content_parts = []
        else:
            if current_id:
                content_parts.append(para.text)
            else:
                preamble_parts.append(para.text)

    if current_id:
        sections.append({
            'id': current_id,
            'heading': current_heading,
            'para_start': current_heading_para,
            'content_para_start': current_heading_para + 1,
            'content_para_end': len(doc.paragraphs),
            'content': '\n'.join(content_parts).strip(),
        })
    elif preamble_parts and not sections:
        sections.append({
            'id': 'full_resume',
            'heading': '',
            'para_start': 0,
            'content_para_start': 0,
            'content_para_end': len(doc.paragraphs),
            'content': '\n'.join(preamble_parts).strip(),
        })

    return sections
