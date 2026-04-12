from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import logging
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
import sqlite3
from datetime import datetime
import io
import difflib
import tempfile

from resume_parser import extract_text_from_pdf, extract_text_from_docx
from ats_matcher import calculate_similarity, find_missing_keywords, ml_score, reset_artifacts
from ai_feedback import generate_ai_suggestions
from ml_train import train_ats_model
from sentence_rewriter import rewrite_sentences
from resume_ner import extract_entities
from smart_update.section_parser import parse_sections_from_text, parse_sections_from_docx
from smart_update.openai_rewriter import rewrite_sections_with_openai, rewrite_single_section, apply_suggestions_with_openai
from smart_update.section_mapper import map_sections, compute_section_diff
from smart_update.docx_preserver import update_docx_in_place
from smart_update.pdf_preserver import generate_updated_pdf

app = Flask(__name__)
CORS(app)

# Make uploads directory absolute to avoid cwd issues when running the server
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Simple SQLite database for auth + per-user history
PROJECT_DATA_DIR = os.path.join(BASE_DIR, 'data')
os.makedirs(PROJECT_DATA_DIR, exist_ok=True)
APPDATA_ROOT = (
    os.environ.get('LOCALAPPDATA')
    or os.environ.get('APPDATA')
    or os.path.join(os.path.expanduser('~'), 'AppData', 'Local')
)
TEMP_APPDATA_ROOT = os.path.join(tempfile.gettempdir(), 'AIResumeAnalyzer')


def _pick_writable_data_dir():
    candidates = [
        os.path.join(APPDATA_ROOT, 'AIResumeAnalyzer'),
        PROJECT_DATA_DIR,
        TEMP_APPDATA_ROOT,
    ]

    for candidate in candidates:
        try:
            os.makedirs(candidate, exist_ok=True)
            probe_path = os.path.join(candidate, 'write_probe.tmp')
            with open(probe_path, 'w', encoding='utf-8') as probe:
                probe.write('ok')
            os.remove(probe_path)
            return candidate
        except OSError:
            continue

    os.makedirs(TEMP_APPDATA_ROOT, exist_ok=True)
    return TEMP_APPDATA_ROOT


DATA_DIR = _pick_writable_data_dir()
DB_PATH = os.path.join(DATA_DIR, 'app.db')
FALLBACK_DB_PATH = os.path.join(DATA_DIR, 'app_recovered.db')

app.config['SECRET_KEY'] = app.config.get('SECRET_KEY') or os.environ.get('RESUME_APP_SECRET', 'dev-secret-change-me')


def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    return conn


def _create_schema(conn):
    with conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT,
                email TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                created_at TEXT NOT NULL,
                last_login TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS analyses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                file_name TEXT,
                ats REAL,
                ml REAL,
                ml_label TEXT,
                jd_snippet TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
            )
            """
        )


def _backup_broken_db():
    if not os.path.exists(DB_PATH):
        return
    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    broken_path = f"{DB_PATH}.broken.{timestamp}"
    try:
        os.replace(DB_PATH, broken_path)
        logger.warning('Existing SQLite database was moved to %s', broken_path)
    except OSError:
        logger.exception('Failed to back up broken database file')


def _switch_to_fallback_db():
    global DB_PATH
    if DB_PATH != FALLBACK_DB_PATH:
        logger.warning('Switching SQLite storage to fallback database %s', FALLBACK_DB_PATH)
        DB_PATH = FALLBACK_DB_PATH


def _is_recoverable_db_error(error: sqlite3.Error) -> bool:
    message = str(error).lower()
    recoverable_markers = (
        'disk i/o error',
        'database disk image is malformed',
        'readonly database',
    )
    return any(marker in message for marker in recoverable_markers)


def _recover_database():
    _backup_broken_db()
    try:
        init_db()
    except sqlite3.Error:
        _switch_to_fallback_db()
        init_db()


def run_db_task(task, retry_on_recoverable: bool = True):
    conn = get_db()
    try:
        return task(conn)
    except sqlite3.Error as db_err:
        if retry_on_recoverable and _is_recoverable_db_error(db_err):
            logger.warning('Recovering SQLite database after error: %s', db_err)
            try:
                conn.close()
            except sqlite3.Error:
                pass
            _recover_database()
            conn = get_db()
            return task(conn)
        raise
    finally:
        try:
            conn.close()
        except sqlite3.Error:
            pass


def init_db():
    conn = None
    try:
        conn = get_db()
        _create_schema(conn)
    except sqlite3.Error as db_err:
        logger.exception('Database initialization failed: %s', db_err)
        if conn is not None:
            try:
                conn.close()
            except sqlite3.Error:
                pass
        _backup_broken_db()
        try:
            conn = get_db()
            _create_schema(conn)
        except sqlite3.Error:
            _switch_to_fallback_db()
            conn = get_db()
            _create_schema(conn)
    finally:
        if conn is not None:
            conn.close()


def _serializer():
    return URLSafeTimedSerializer(app.config['SECRET_KEY'], salt='auth-token')


def create_token(user_id: int) -> str:
    return _serializer().dumps({'user_id': user_id})


def get_current_user_id():
    """Return user_id from Authorization: Bearer <token> header, or None."""
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '):
        return None
    token = auth_header.split(' ', 1)[1].strip()
    if not token:
        return None
    try:
        data = _serializer().loads(token, max_age=60 * 60 * 24 * 60)  # 60 days
        return int(data.get('user_id'))
    except (BadSignature, SignatureExpired, ValueError, TypeError):
        return None


# Basic logging for server-side diagnostics
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure DB exists on startup
init_db()

# 🏠 Route: Home
@app.route('/')
def home():
    return "✅ Resume Analyzer API is running!"

# Return JSON for any uncaught exceptions to aid debugging in frontend Network tab
@app.errorhandler(Exception)
def handle_exception(e):
    logger.exception('Unhandled exception: %s', e)
    return jsonify({'error': 'Internal server error', 'detail': str(e)}), 500

# 📤 Route 1: Upload Resume and Extract Text
@app.route('/upload', methods=['POST'])
def upload_resume():
    try:
        if 'resume' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 200

        file = request.files['resume']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 200

        safe_name = secure_filename(file.filename)
        if not safe_name:
            return jsonify({'success': False, 'error': 'Invalid file name'}), 200

        filepath = os.path.join(UPLOAD_FOLDER, safe_name)
        try:
            file.save(filepath)
        except Exception as save_err:
            logger.exception('Failed to save upload: %s', save_err)
            return jsonify({'success': False, 'error': 'Failed to save file on server'}), 200

        filename_lower = safe_name.lower()
        text = ''
        warning = None

        try:
            if filename_lower.endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename_lower.endswith('.docx'):
                text = extract_text_from_docx(filepath)
            else:
                return jsonify({'success': False, 'error': 'Unsupported file type'}), 200
        except Exception as parse_err:
            logger.exception('Failed to parse uploaded file: %s', parse_err)
            # Avoid 500 on tricky PDFs/Docs; return empty text with a warning instead
            warning = 'File uploaded but text extraction failed. Try another file format.'

        return jsonify({'success': True, 'resume_text': (text or ''), 'warning': warning})
    except Exception as err:
        logger.exception('Unexpected error in /upload: %s', err)
        return jsonify({'success': False, 'error': 'Unexpected server error'}), 200


def _save_resume_upload(file):
    if file is None:
        raise ValueError('Resume file is required')

    safe_name = secure_filename(file.filename or '')
    if not safe_name:
        raise ValueError('Invalid file name')

    filepath = os.path.join(UPLOAD_FOLDER, safe_name)
    file.save(filepath)
    return safe_name, filepath


def _read_resume_text(file_name, filepath):
    lower_name = (file_name or '').lower()
    if lower_name.endswith('.pdf'):
        return extract_text_from_pdf(filepath)
    if lower_name.endswith('.docx'):
        return extract_text_from_docx(filepath)
    raise ValueError('Unsupported file type')

# 🧠 Route 2: Analyze Resume vs JD (ATS + ML + Keywords)
@app.route('/analyze', methods=['POST'])
def analyze_resume():
    if 'resume' not in request.files or 'jd' not in request.form:
        return jsonify({'error': 'Missing resume or job description'}), 400

    file = request.files['resume']
    jd_text = request.form['jd']

    try:
        safe_name, filepath = _save_resume_upload(file)
        resume_text = _read_resume_text(safe_name, filepath)
    except ValueError as err:
        return jsonify({'error': str(err)}), 400
    except Exception as err:
        logger.exception('Analyze route failed before scoring: %s', err)
        return jsonify({'error': 'Failed to process the uploaded resume'}), 500

    score = calculate_similarity(resume_text, jd_text)
    keywords = find_missing_keywords(resume_text, jd_text)
    ml_result = ml_score(resume_text, jd_text)

    # If the request is authenticated, log this analysis for the user
    try:
        user_id = get_current_user_id()
        if user_id is not None:
            conn = get_db()
            with conn:
                conn.execute(
                    """
                    INSERT INTO analyses (user_id, file_name, ats, ml, ml_label, jd_snippet, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        user_id,
                        safe_name,
                        float(score or 0.0),
                        float(ml_result.get('ml_score') or 0.0),
                        ml_result.get('ml_label') or '',
                        (jd_text or '')[:200],
                        datetime.utcnow().isoformat(timespec='seconds'),
                    ),
                )
    except Exception as log_err:
        logger.exception('Failed to log analysis: %s', log_err)

    return jsonify({
        'ats_score': score,
        'missing_keywords': keywords['missing_keywords'],
        'matched_keywords': keywords['matched_keywords'],
        'ml_score': ml_result.get('ml_score'),
        'ml_label': ml_result.get('ml_label'),
        'ml_details': ml_result.get('details', {})
    })

# 🤖 Route 3: AI Suggestions (NLP + NER + Grammar)
@app.route('/ai-suggest', methods=['POST'])
def ai_suggest():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required'}), 400

    file = request.files['resume']
    jd_text = request.form.get('jd', '')

    try:
        safe_name, filepath = _save_resume_upload(file)
        resume_text = _read_resume_text(safe_name, filepath)
    except ValueError as err:
        return jsonify({'error': str(err)}), 400
    except Exception as err:
        logger.exception('AI suggestions failed: %s', err)
        return jsonify({'error': 'Failed to process the uploaded resume'}), 500

    suggestions = generate_ai_suggestions(resume_text, jd_text)
    return jsonify({'ai_feedback': suggestions})

# 📊 Route 4: NER Extraction Endpoint
@app.route('/extract-ner', methods=['POST'])
def extract_ner():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required'}), 400

    try:
        safe_name, filepath = _save_resume_upload(request.files['resume'])
        resume_text = _read_resume_text(safe_name, filepath)
    except ValueError as err:
        return jsonify({'error': str(err)}), 400
    except Exception as err:
        logger.exception('NER extraction failed: %s', err)
        return jsonify({'error': 'Failed to process the uploaded resume'}), 500

    ner_data = extract_entities(resume_text)
    return jsonify(ner_data)

# ✍️ Route 5: Dynamic Resume Sentence Rewriting
@app.route('/rewrite', methods=['POST'])
def rewrite_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required'}), 400

    try:
        safe_name, filepath = _save_resume_upload(request.files['resume'])
        resume_text = _read_resume_text(safe_name, filepath)
    except ValueError as err:
        return jsonify({'error': str(err)}), 400
    except Exception as err:
        logger.exception('Rewrite route failed: %s', err)
        return jsonify({'error': 'Failed to process the uploaded resume'}), 500

    rewritten = rewrite_sentences(resume_text)
    return jsonify(rewritten)


# 🖍️ Route 6: Generate HTML diff preview for highlighted changes
@app.route('/preview-diff', methods=['POST'])
def preview_diff():
    data = request.get_json(silent=True) or {}
    original_text = data.get('original_text', '')
    updated_text = data.get('updated_text', '')
    if not isinstance(original_text, str) or not isinstance(updated_text, str):
        return jsonify({'error': 'original_text and updated_text must be strings'}), 400

    # Use difflib to build an inline HTML diff
    differ = difflib.HtmlDiff(wrapcolumn=80)
    html = differ.make_table(
        original_text.splitlines(),
        updated_text.splitlines(),
        fromdesc='Original',
        todesc='Updated',
        context=True,
        numlines=2
    )
    # Wrap in minimal HTML so frontend can directly render
    doc = f"""
<!doctype html>
<html>
<head>
  <meta charset='utf-8'>
  <style>
    body {{ font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; }}
    table.diff {{ font-size: 14px; border-collapse: collapse; width: 100%; }}
    .diff_header {{ background: #f6f7f9; }}
    .diff_add {{ background: #e6ffed; }}
    .diff_chg {{ background: #fff5b1; }}
    .diff_sub {{ background: #ffeef0; }}
    td, th {{ border: 1px solid #dfe2e5; padding: 4px 6px; vertical-align: top; }}
  </style>
  <title>Resume Changes Preview</title>
</head>
<body>
{html}
</body>
</html>
"""
    return jsonify({'html': doc})


# ⬇️ Route 7: Download updated resume as PDF or DOCX
@app.route('/download-updated', methods=['POST'])
def download_updated():
    fmt = request.args.get('format', 'pdf').lower()

    # Accept either JSON body or multipart (for future extensibility)
    text = ''
    if request.is_json:
        text = (request.get_json(silent=True) or {}).get('updated_text', '')
    else:
        text = request.form.get('updated_text', '')

    if not isinstance(text, str) or not text.strip():
        return jsonify({'error': 'updated_text is required'}), 400

    if fmt == 'docx':
        # Generate DOCX using python-docx (already in requirements)
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            return jsonify({'error': f'DOCX generation unavailable: {e}'}), 500

        doc = Document()
        for para in text.split('\n\n'):
            for line in para.split('\n'):
                doc.add_paragraph(line)
            doc.add_paragraph('')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name='resume_updated.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    # Default: PDF using reportlab if available; fallback to plain text PDF-like
    try:
        from reportlab.lib.pagesizes import LETTER  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
    except Exception as e:
        return jsonify({'error': f'PDF generation unavailable: {e}'}), 500

    # Prefer Platypus for safer wrapping; fallback to low-level canvas
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
        styles = getSampleStyleSheet()
        normal = styles['Normal']

        # sanitize: replace problematic unicode, and escape HTML-ish chars
        def esc(s: str) -> str:
            s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            s = s.replace('\u2022', '-').replace('•', '-')
            s = s.replace('\u2013', '-').replace('\u2014', '-').replace('–', '-').replace('—', '-')
            try:
                s.encode('latin-1')
            except Exception:
                s = s.encode('ascii', 'ignore').decode('ascii')
            return s

        story = []
        for para in text.split('\n\n'):
            p = Paragraph(esc(para).replace('\n', '<br/>'), normal)
            story.append(p)
            story.append(Spacer(1, 6))
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=LETTER)
        doc.build(story)
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name='resume_updated.pdf', mimetype='application/pdf')
    except Exception:
        # Fallback to manual canvas rendering
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=LETTER)
        width, height = LETTER
        left_margin = 1 * inch
        top_margin = height - 1 * inch
        line_height = 12

        def _sanitize_for_pdf(s: str) -> str:
            try:
                s = s.replace('\u2022', '-').replace('\u2013', '-').replace('\u2014', '-')
                s = s.replace('•', '-').replace('–', '-').replace('—', '-')
                return s.encode('ascii', 'ignore').decode('ascii')
            except Exception:
                return s

        y = top_margin
        for raw_line in text.split('\n'):
            line = raw_line.replace('\t', '    ')
            line = _sanitize_for_pdf(line)
            words = line.split(' ')
            current = ''
            for w in words:
                test = (current + ' ' + w).strip()
                if c.stringWidth(test, 'Helvetica', 10) > (width - 2 * left_margin):
                    c.setFont('Helvetica', 10)
                    c.drawString(left_margin, y, current)
                    y -= line_height
                    current = w
                    if y < 1 * inch:
                        c.showPage()
                        y = top_margin
                else:
                    current = test
            if current:
                c.setFont('Helvetica', 10)
                c.drawString(left_margin, y, current)
                y -= line_height
                if y < 1 * inch:
                    c.showPage()
                    y = top_margin
        c.save()
        buf.seek(0)
        return send_file(buf, as_attachment=True, download_name='resume_updated.pdf', mimetype='application/pdf')


# ─── Smart Resume Update Routes ────────────────────────────────────────────

# Health check endpoint for frontend connectivity detection
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok'}), 200


# 📋 Smart Parse: Extract structured sections from resume
@app.route('/smart-parse', methods=['POST'])
def smart_parse():
    if 'resume' not in request.files:
        return jsonify({'error': 'Resume file is required'}), 400

    file = request.files['resume']
    safe_name = secure_filename(file.filename)
    if not safe_name:
        return jsonify({'error': 'Invalid file name'}), 400

    try:
        _, filepath = _save_resume_upload(file)
    except ValueError as err:
        return jsonify({'error': str(err)}), 400
    except Exception as err:
        logger.exception('Smart parse save failed: %s', err)
        return jsonify({'error': 'Failed to save the uploaded resume'}), 500

    filename_lower = safe_name.lower()
    try:
        if filename_lower.endswith('.docx'):
            sections = parse_sections_from_docx(filepath)
            # Also extract full text for reference
            full_text = extract_text_from_docx(filepath)
            file_type = 'docx'
        elif filename_lower.endswith('.pdf'):
            full_text = extract_text_from_pdf(filepath)
            sections = parse_sections_from_text(full_text)
            file_type = 'pdf'
        else:
            return jsonify({'error': 'Unsupported file type. Use .pdf or .docx'}), 400
    except Exception as e:
        logger.exception('Smart parse failed: %s', e)
        return jsonify({'error': f'Failed to parse resume: {str(e)}'}), 500

    return jsonify({
        'success': True,
        'file_type': file_type,
        'file_name': safe_name,
        'full_text': full_text,
        'sections': sections,
        'section_count': len(sections),
    })


# 🤖 Smart Rewrite: Send sections to OpenAI for improvement
@app.route('/smart-rewrite', methods=['POST'])
def smart_rewrite():
    data = request.get_json(silent=True) or {}
    sections = data.get('sections', [])
    job_description = data.get('job_description', '')
    model = data.get('model', 'gpt-4o-mini')

    if not sections:
        return jsonify({'error': 'No sections provided'}), 400

    try:
        ai_results = rewrite_sections_with_openai(sections, job_description, model)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except ImportError as e:
        return jsonify({'error': str(e)}), 500
    except Exception as e:
        logger.exception('OpenAI rewrite failed: %s', e)
        return jsonify({'error': f'AI rewrite failed: {str(e)}'}), 500

    # Map AI results back to original sections
    mappings = map_sections(sections, ai_results)

    # Compute diffs for each section
    for m in mappings:
        if m.get('has_changes'):
            m['diff'] = compute_section_diff(
                m.get('original_content', ''),
                m.get('improved_content', '')
            )

    return jsonify({
        'success': True,
        'mappings': mappings,
        'ai_sections': ai_results,
        'sections_improved': sum(1 for m in mappings if m.get('has_changes')),
    })


# 🔄 Smart Rewrite Single: Rewrite one section at a time
@app.route('/smart-rewrite-section', methods=['POST'])
def smart_rewrite_section():
    data = request.get_json(silent=True) or {}
    section_id = data.get('section_id', '')
    section_content = data.get('content', '')
    section_heading = data.get('heading', '')
    job_description = data.get('job_description', '')

    if not section_id or not section_content:
        return jsonify({'error': 'section_id and content are required'}), 400

    try:
        improved = rewrite_single_section(
            section_id, section_content, section_heading, job_description
        )
    except Exception as e:
        logger.exception('Section rewrite failed: %s', e)
        return jsonify({'error': f'Section rewrite failed: {str(e)}'}), 500

    diff = compute_section_diff(section_content, improved)

    return jsonify({
        'success': True,
        'section_id': section_id,
        'original_content': section_content,
        'improved_content': improved,
        'has_changes': improved.strip() != section_content.strip(),
        'diff': diff,
    })


# 🎯 Apply Suggestions: Take AI suggestions and apply them to resume preserving formatting
@app.route('/apply-suggestions', methods=['POST'])
def apply_suggestions():
    data = request.get_json(silent=True) or {}
    sections = data.get('sections', [])
    suggestions = data.get('suggestions', '')
    job_description = data.get('job_description', '')
    model = data.get('model', 'gpt-4o-mini')

    if not sections:
        return jsonify({'error': 'No sections provided'}), 400
    if not suggestions or not suggestions.strip():
        return jsonify({'error': 'No suggestions provided'}), 400

    try:
        ai_results = apply_suggestions_with_openai(sections, suggestions, job_description, model)
    except Exception as e:
        logger.exception('Apply suggestions failed: %s', e)
        return jsonify({'error': f'Failed to apply suggestions: {str(e)}'}), 500

    if not ai_results:
        return jsonify({
            'success': True,
            'mappings': [],
            'sections_improved': 0,
            'message': 'No applicable changes were generated from the suggestions.',
        })

    # Map AI results back to original sections (reuse existing mapping logic)
    mappings = map_sections(sections, ai_results)

    # Compute diffs for each section
    for m in mappings:
        if m.get('has_changes'):
            m['diff'] = compute_section_diff(
                m.get('original_content', ''),
                m.get('improved_content', '')
            )

    return jsonify({
        'success': True,
        'mappings': mappings,
        'sections_improved': sum(1 for m in mappings if m.get('has_changes')),
    })


# 📥 Smart Download: Generate updated file preserving original formatting
@app.route('/smart-download', methods=['POST'])
def smart_download():
    data = request.get_json(silent=True) or {}
    file_name = data.get('file_name', '')
    file_type = data.get('file_type', 'pdf')
    mappings = data.get('mappings', [])
    accepted_sections = data.get('accepted_sections')  # None = accept all
    output_format = data.get('output_format', file_type)  # default to original format

    if not mappings:
        return jsonify({'error': 'No section mappings provided'}), 400

    try:
        if output_format == 'docx' and file_name:
            filepath = os.path.join(UPLOAD_FOLDER, file_name)
            if os.path.exists(filepath) and filepath.lower().endswith('.docx'):
                # Use format-preserving DOCX update
                buf = update_docx_in_place(filepath, mappings, accepted_sections)
                return send_file(
                    buf,
                    as_attachment=True,
                    download_name='resume_improved.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

        # For PDF output or when original DOCX isn't available
        # Reconstruct full text from mappings
        full_text = data.get('original_text', '')
        buf = generate_updated_pdf(full_text, mappings, accepted_sections)
        return send_file(
            buf,
            as_attachment=True,
            download_name='resume_improved.pdf',
            mimetype='application/pdf'
        )

    except Exception as e:
        logger.exception('Smart download failed: %s', e)
        return jsonify({'error': f'Download failed: {str(e)}'}), 500


# 📋 Smart Preview: Get section-by-section diff without downloading
@app.route('/smart-preview', methods=['POST'])
def smart_preview():
    data = request.get_json(silent=True) or {}
    mappings = data.get('mappings', [])
    accepted_sections = data.get('accepted_sections')

    if not mappings:
        return jsonify({'error': 'No section mappings provided'}), 400

    # Build original and updated full text for HTML diff
    original_parts = []
    updated_parts = []

    for m in mappings:
        heading = m.get('original_heading', '')
        original_content = m.get('original_content', '')
        improved_content = m.get('improved_content', original_content)
        section_id = m.get('original_id', '')

        use_improved = (
            m.get('has_changes', False)
            and (accepted_sections is None or section_id in accepted_sections)
        )

        if heading:
            original_parts.append(heading)
            updated_parts.append(heading)

        original_parts.append(original_content)
        updated_parts.append(improved_content if use_improved else original_content)
        original_parts.append('')
        updated_parts.append('')

    original_text = '\n'.join(original_parts)
    updated_text = '\n'.join(updated_parts)

    # Generate HTML diff
    differ = difflib.HtmlDiff(wrapcolumn=80)
    html = differ.make_table(
        original_text.splitlines(),
        updated_text.splitlines(),
        fromdesc='Original',
        todesc='Improved',
        context=True,
        numlines=2
    )
    doc = f"""
<!doctype html>
<html>
<head>
  <meta charset='utf-8'>
  <style>
    body {{ font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; }}
    table.diff {{ font-size: 14px; border-collapse: collapse; width: 100%; }}
    .diff_header {{ background: #f6f7f9; }}
    .diff_add {{ background: #e6ffed; }}
    .diff_chg {{ background: #fff5b1; }}
    .diff_sub {{ background: #ffeef0; }}
    td, th {{ border: 1px solid #dfe2e5; padding: 4px 6px; vertical-align: top; }}
  </style>
</head>
<body>{html}</body>
</html>
"""
    return jsonify({
        'html': doc,
        'original_text': original_text,
        'updated_text': updated_text,
    })


# ─── End Smart Resume Update Routes ───────────────────────────────────────


# 👤 Auth: Sign up
@app.route('/auth/signup', methods=['POST'])
def auth_signup():
    data = request.get_json(silent=True) or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    if len(password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters'}), 400

    now = datetime.utcnow().isoformat(timespec='seconds')

    def _create_user(conn):
        with conn:
            cur = conn.execute('SELECT id FROM users WHERE email = ?', (email,))
            if cur.fetchone():
                return 'exists'
            cur = conn.execute(
                'INSERT INTO users (name, email, password_hash, created_at, last_login) VALUES (?, ?, ?, ?, ?)',
                (name or None, email, generate_password_hash(password), now, now),
            )
            return cur.lastrowid

    try:
        user_id = run_db_task(_create_user)
        if user_id == 'exists':
            return jsonify({'error': 'An account with this email already exists'}), 400
    except sqlite3.Error as db_err:
        logger.exception('Failed to create user: %s', db_err)
        return jsonify({'error': 'Failed to create account'}), 500

    token = create_token(user_id)
    return jsonify({
        'token': token,
        'user': {
            'id': user_id,
            'name': name,
            'email': email,
            'created_at': now,
        },
    })


# 👤 Auth: Login
@app.route('/auth/login', methods=['POST'])
def auth_login():
    data = request.get_json(silent=True) or {}
    email = (data.get('email') or '').strip().lower()
    password = data.get('password') or ''

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    now = datetime.utcnow().isoformat(timespec='seconds')

    def _login(conn):
        cur = conn.execute('SELECT id, name, email, password_hash, created_at FROM users WHERE email = ?', (email,))
        row = cur.fetchone()
        if not row or not check_password_hash(row['password_hash'], password):
            return None

        with conn:
            conn.execute('UPDATE users SET last_login = ? WHERE id = ?', (now, row['id']))
        return row

    try:
        row = run_db_task(_login)
        if not row:
            return jsonify({'error': 'Invalid email or password'}), 401
    except sqlite3.Error as db_err:
        logger.exception('Login failed: %s', db_err)
        return jsonify({'error': 'Login failed'}), 500

    token = create_token(row['id'])
    return jsonify({
        'token': token,
        'user': {
            'id': row['id'],
            'name': row['name'],
            'email': row['email'],
            'created_at': row['created_at'],
        },
    })


@app.route('/auth/me', methods=['GET'])
def auth_me():
    user_id = get_current_user_id()
    if user_id is None:
        return jsonify({'error': 'Authentication required'}), 401

    def _load_user(conn):
        cur = conn.execute(
            'SELECT id, name, email, created_at, last_login FROM users WHERE id = ?',
            (user_id,),
        )
        return cur.fetchone()

    try:
        row = run_db_task(_load_user)
        if row is None:
            return jsonify({'error': 'User not found'}), 404
    except sqlite3.Error as db_err:
        logger.exception('Failed to load current user: %s', db_err)
        return jsonify({'error': 'Failed to load current user'}), 500

    return jsonify({
        'user': {
            'id': row['id'],
            'name': row['name'],
            'email': row['email'],
            'created_at': row['created_at'],
            'last_login': row['last_login'],
        }
    })


# 📜 Per-user analysis history
@app.route('/history', methods=['GET'])
def get_history():
    user_id = get_current_user_id()
    if user_id is None:
        return jsonify({'error': 'Authentication required'}), 401

    def _load_history(conn):
        cur = conn.execute(
            """
            SELECT id, file_name, ats, ml, ml_label, jd_snippet, created_at
            FROM analyses
            WHERE user_id = ?
            ORDER BY datetime(created_at) DESC
            LIMIT 20
            """,
            (user_id,),
        )
        return [
            {
                'id': r['id'],
                'file_name': r['file_name'],
                'ats': r['ats'],
                'ml': r['ml'],
                'ml_label': r['ml_label'],
                'jd_snippet': r['jd_snippet'],
                'created_at': r['created_at'],
            }
            for r in cur.fetchall()
        ]

    try:
        rows = run_db_task(_load_history)
    except sqlite3.Error as db_err:
        logger.exception('Failed to load history: %s', db_err)
        return jsonify({'error': 'Failed to load history'}), 500

    return jsonify({'items': rows})


@app.route('/admin/retrain-model', methods=['POST'])
def retrain_model():
    try:
        summary = train_ats_model()
        reset_artifacts()
        return jsonify({
            'success': True,
            'summary': summary,
        })
    except Exception as err:
        logger.exception('Model retraining failed: %s', err)
        return jsonify({'error': f'Model retraining failed: {str(err)}'}), 500


# 🔚 App runner — always keep this LAST
if __name__ == '__main__':
    app.run(debug=True)
